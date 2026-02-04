"""文档加载器 - 支持多种格式"""

import logging
import tempfile
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
import markdown

from ..config import settings
from .image_encoder import image_encoder

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """加载后的文档"""
    content: str
    metadata: dict
    images: List[dict] = field(default_factory=list)  # [{"data": bytes, "page": int, "name": str}, ...]
    

class DocumentLoader:
    """多格式文档加载器"""
    
    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "word",
        ".doc": "word",
        ".xlsx": "excel",
        ".xls": "excel",
        ".md": "markdown",
        ".txt": "text",
    }
    
    @classmethod
    def get_file_type(cls, filename: str) -> Optional[str]:
        """获取文件类型"""
        suffix = Path(filename).suffix.lower()
        return cls.SUPPORTED_TYPES.get(suffix)
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查文件是否支持"""
        return cls.get_file_type(filename) is not None
    
    def load(self, file_path: str, filename: str) -> LoadedDocument:
        """加载文档"""
        file_type = self.get_file_type(filename)
        
        if file_type == "pdf":
            return self._load_pdf(file_path, filename)
        elif file_type == "word":
            return self._load_word(file_path, filename)
        elif file_type == "excel":
            return self._load_excel(file_path, filename)
        elif file_type == "markdown":
            return self._load_markdown(file_path, filename)
        elif file_type == "text":
            return self._load_text(file_path, filename)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    
    def _load_pdf(self, file_path: str, filename: str) -> LoadedDocument:
        """加载PDF文档"""
        logger.info(f"Loading PDF: {filename}")

        mode = (getattr(settings, "pdf_extractor", "marker") or "marker").strip().lower()

        if mode in ("marker", "auto"):
            # Prefer Marker for better structure.
            # Important: for中文 PDF，不要用默认的 ocr_alphanum_threshold=0.3；否则容易误判为乱码导致全量 OCR。
            try:
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                from marker.renderers.chunk import ChunkOutput

                def _html_to_text(html: str) -> str:
                    """Convert marker block HTML to readable text.

                    Important: keep inline math structure (e.g. <sup>/<sub>) without inserting newlines.
                    """
                    try:
                        from bs4 import BeautifulSoup
                        from bs4.element import NavigableString, Tag
                        import re

                        soup = BeautifulSoup(str(html or ""), "html.parser")

                        def inline(node) -> str:
                            if node is None:
                                return ""
                            if isinstance(node, NavigableString):
                                return str(node)
                            if not isinstance(node, Tag):
                                return str(node)

                            name = (node.name or "").lower()
                            if name == "sup":
                                inner = "".join(inline(c) for c in node.children).strip()
                                return f"^{{{inner}}}" if inner else ""
                            if name == "sub":
                                inner = "".join(inline(c) for c in node.children).strip()
                                return f"_{{{inner}}}" if inner else ""
                            if name == "br":
                                return "\n"
                            if name == "math":
                                inner = "".join(inline(c) for c in node.children).strip()
                                if not inner:
                                    inner = node.get_text("", strip=True)
                                if not inner:
                                    return ""
                                return f"\n$$\n{inner}\n$$\n"
                            if name in ("del", "s", "strike"):
                                # Drop deleted content.
                                return ""

                            return "".join(inline(c) for c in node.children)

                        def _normalize_text(s: str) -> str:
                            x = str(s or "")
                            # Normalize common dash variants that appear in OCR/PDF extraction.
                            x = x.replace("−", "-")  # minus sign
                            x = x.replace("—", "-")  # em dash
                            x = re.sub(r"[ \t]+", " ", x)
                            return x

                        def table_to_markdown(table: Tag) -> str:
                            rows = []
                            for tr in table.find_all("tr"):
                                cells = []
                                for cell in tr.find_all(["th", "td"]):
                                    t = inline(cell)
                                    t = _normalize_text(t)
                                    # Collapse whitespace/newlines inside table cells.
                                    t = re.sub(r"\s*\n\s*", "<br/>", t.strip())
                                    # Convert display math blocks to inline math inside table cells.
                                    t = re.sub(
                                        r"\$\$\s*([\s\S]*?)\s*\$\$",
                                        lambda m: "$" + re.sub(r"\s+", " ", m.group(1)).strip() + "$",
                                        t,
                                    )
                                    t = t.replace("|", "\\|")
                                    cells.append(t)
                                if cells:
                                    rows.append(cells)

                            if not rows:
                                return ""

                            col_count = max(len(r) for r in rows)
                            rows = [r + [""] * (col_count - len(r)) for r in rows]

                            header = rows[0]
                            sep = ["---"] * col_count
                            body = rows[1:] if len(rows) > 1 else []

                            # If this is a formula table, render the formula column as math.
                            formula_col = None
                            for i, h in enumerate(header):
                                hh = str(h or "")
                                hh = hh.replace("<br/>", " ")
                                hh = re.sub(r"\s+", "", hh)
                                if "计算公式" in hh:
                                    formula_col = i
                                    break

                            if formula_col is not None:
                                def ensure_math(cell_text: str) -> str:
                                    s = str(cell_text or "").strip()
                                    if not s:
                                        return s
                                    # already has math delimiters
                                    if "$" in s:
                                        # Normalize patterns like: i = $\frac{...}{...}$ -> $i = \frac{...}{...}$
                                        m = re.match(r"^\s*([A-Za-z][A-Za-z0-9_]*\s*=\s*)\$\s*([^$]+?)\s*\$\s*$", s)
                                        if m:
                                            return "$" + m.group(1).strip() + m.group(2).strip() + "$"

                                        # fix unbalanced $ in-cell
                                        if s.count("$") % 2 == 1:
                                            s += "$"
                                        return s

                                    # If the cell contains multiple lines, wrap each part.
                                    parts = [p.strip() for p in s.split("<br/>")]
                                    wrapped_parts = []
                                    for p in parts:
                                        if not p:
                                            continue
                                        wrapped_parts.append(f"${p}$")
                                    return "<br/>".join(wrapped_parts) if wrapped_parts else s

                                for r in body:
                                    if formula_col < len(r):
                                        r[formula_col] = ensure_math(r[formula_col])
                            else:
                                # Fallback: detect a likely formula column by content.
                                # This keeps behavior general while ensuring tables with TeX fragments are renderable.
                                def score_cell(s: str) -> int:
                                    t = str(s or "")
                                    sc = 0
                                    if "$" in t:
                                        sc += 3
                                    if "\\frac" in t or "\\sum" in t or "\\left" in t or "\\right" in t:
                                        sc += 3
                                    if "^{" in t or "_{" in t:
                                        sc += 2
                                    if "=" in t:
                                        sc += 1
                                    return sc

                                best_i = None
                                best_score = 0
                                for ci in range(col_count):
                                    ssum = 0
                                    for r in body:
                                        if ci < len(r):
                                            ssum += score_cell(r[ci])
                                    if ssum > best_score:
                                        best_score = ssum
                                        best_i = ci

                                if best_i is not None and best_score >= 6:
                                    def ensure_math(cell_text: str) -> str:
                                        s = str(cell_text or "").strip()
                                        if not s:
                                            return s
                                        if "$" in s:
                                            if s.count("$") % 2 == 1:
                                                s += "$"
                                            return s
                                        parts = [p.strip() for p in s.split("<br/>")]
                                        wrapped_parts = []
                                        for p in parts:
                                            if not p:
                                                continue
                                            wrapped_parts.append(f"${p}$")
                                        return "<br/>".join(wrapped_parts) if wrapped_parts else s

                                    for r in body:
                                        if best_i < len(r):
                                            r[best_i] = ensure_math(r[best_i])

                            def fmt(r):
                                return "| " + " | ".join(r) + " |"

                            out = [fmt(header), fmt(sep)]
                            out.extend(fmt(r) for r in body)
                            return "\n".join(out).strip()

                        lines: list[str] = []

                        # Tables: keep structure as markdown table.
                        tbl = soup.find("table")
                        if isinstance(tbl, Tag):
                            md = table_to_markdown(tbl)
                            if md:
                                return md

                        for li in soup.find_all("li"):
                            t = inline(li).strip()
                            if t:
                                lines.append(f"- {t}")
                        if lines:
                            txt = "\n".join(lines)
                        else:
                            ps = soup.find_all("p")
                            if ps:
                                parts = []
                                for p in ps:
                                    t = inline(p).strip()
                                    if t:
                                        parts.append(t)
                                txt = "\n\n".join(parts)
                            else:
                                txt = inline(soup).strip()

                        # Normalize whitespace while preserving intentional newlines.
                        out_lines = []
                        for ln in str(txt).splitlines():
                            out_lines.append(_normalize_text(ln).rstrip())
                        out = "\n".join(out_lines)
                        out = "\n".join([ln for ln in out.splitlines() if ln.strip() or ln == ""])

                        # Wrap bare TeX-like formula lines into math mode.
                        # This avoids leaking raw `\frac/\times/...` into the UI when marker didn't emit <math>.
                        fixed_lines = []
                        in_display_math = False
                        for ln in out.splitlines():
                            t = ln.strip()
                            if not t:
                                fixed_lines.append(ln)
                                continue
                            if t == "$$":
                                in_display_math = not in_display_math
                                fixed_lines.append(ln)
                                continue
                            if in_display_math:
                                fixed_lines.append(ln)
                                continue
                            if "$" in t:
                                # Normalize patterns like: i = $\frac{...}{...}$ -> $i = \frac{...}{...}$
                                m = re.match(r"^\s*([A-Za-z][A-Za-z0-9_]*\s*=\s*)\$\s*([^$]+?)\s*\$\s*$", t)
                                if m:
                                    fixed_lines.append("$" + m.group(1).strip() + m.group(2).strip() + "$")
                                else:
                                    fixed_lines.append(ln)
                                continue
                            if t.startswith("#") or t.startswith("-") or t.startswith(">"):
                                fixed_lines.append(ln)
                                continue

                            has_tex_cmd = bool(re.search(r"\\[a-zA-Z]+", t))
                            has_ops = "=" in t
                            if has_tex_cmd and has_ops:
                                fixed_lines.append(f"${t}$")
                                continue

                            fixed_lines.append(ln)
                        out = "\n".join(fixed_lines)

                        # Collapse excessive blank lines.
                        while "\n\n\n" in out:
                            out = out.replace("\n\n\n", "\n\n")
                        return out.strip()
                    except Exception:
                        return str(html or "").strip()

                def _build_section_id_title_map(rendered: ChunkOutput) -> dict[str, str]:
                    m: dict[str, str] = {}
                    for b in rendered.blocks or []:
                        if str(getattr(b, "block_type", "")).lower() != "sectionheader":
                            continue
                        sid = str(getattr(b, "id", "") or "").strip()
                        if not sid:
                            continue
                        txt = _html_to_text(getattr(b, "html", ""))
                        if txt:
                            # Prefer the first line as heading text.
                            m[sid] = txt.split("\n", 1)[0].strip()
                    return m

                def _extract_outline_headings_from_chunk_output(rendered: ChunkOutput, id_to_title: dict[str, str]) -> list[tuple[int, str]]:
                    headings: list[tuple[int, str]] = []
                    last_path: list[tuple[int, str]] = []
                    for b in rendered.blocks or []:
                        hierarchy = getattr(b, "section_hierarchy", None) or None
                        if not hierarchy:
                            continue

                        # hierarchy value is a section header block id, not the title.
                        path: list[tuple[int, str]] = []
                        for k, v in hierarchy.items():
                            sid = str(v).strip()
                            if not sid:
                                continue
                            title = (id_to_title.get(sid) or sid).strip()
                            if not title:
                                continue
                            path.append((int(k), title))
                        path.sort(key=lambda x: x[0])

                        i = 0
                        while i < len(last_path) and i < len(path) and last_path[i] == path[i]:
                            i += 1
                        for lvl, title in path[i:]:
                            headings.append((max(1, min(6, int(lvl))), title))
                        last_path = path
                    return headings

                def _build_markdown_from_chunk_output(rendered: ChunkOutput, id_to_title: dict[str, str]) -> str:
                    md_lines: list[str] = []
                    last_path: list[tuple[int, str]] = []
                    for b in rendered.blocks or []:
                        hierarchy = getattr(b, "section_hierarchy", None) or None
                        if hierarchy:
                            path: list[tuple[int, str]] = []
                            for k, v in hierarchy.items():
                                sid = str(v).strip()
                                if not sid:
                                    continue
                                title = (id_to_title.get(sid) or sid).strip()
                                if not title:
                                    continue
                                path.append((int(k), title))
                            path.sort(key=lambda x: x[0])

                            i = 0
                            while i < len(last_path) and i < len(path) and last_path[i] == path[i]:
                                i += 1
                            for lvl, title in path[i:]:
                                md_lines.append("#" * max(1, min(6, int(lvl))) + " " + title)
                                md_lines.append("")
                            last_path = path

                        # Skip section header block body (already emitted as a heading).
                        if str(getattr(b, "block_type", "")).lower() == "sectionheader":
                            continue

                        txt = _html_to_text(getattr(b, "html", ""))
                        if txt:
                            md_lines.append(txt)
                            md_lines.append("")

                    return "\n".join(md_lines).strip()

                converter = PdfConverter(
                    artifact_dict=create_model_dict(),
                    renderer="marker.renderers.chunk.ChunkRenderer",
                    config={
                        "force_ocr": bool(getattr(settings, "pdf_marker_force_ocr", False)),
                        "ocr_alphanum_threshold": float(getattr(settings, "pdf_marker_ocr_alphanum_threshold", 0.0)),
                    },
                )
                rendered = None
                for attempt in range(2):
                    try:
                        rendered = converter(file_path)
                        break
                    except Exception as e:
                        msg = str(e)
                        if attempt == 0 and ("IncompleteRead" in msg or "Connection broken" in msg):
                            continue
                        raise

                if not isinstance(rendered, ChunkOutput):
                    raise ValueError(f"Unexpected marker output type: {type(rendered)}")

                id_to_title = _build_section_id_title_map(rendered)
                md_text = _build_markdown_from_chunk_output(rendered, id_to_title)
                outline_headings = _extract_outline_headings_from_chunk_output(rendered, id_to_title)

                # Merge standalone exponent blocks like:
                #   <expr>\n$$\n^{5}\n$$
                # into:
                #   <expr>^{5}
                try:
                    import re as _re

                    # Case 1: plain line followed by $$ exponent $$
                    md_text = _re.sub(
                        r"(?m)^(?P<base>[^\n$].*?)\n\$\$\n(?P<exp>[\^_]\{[^}]+\})\n\$\$\s*$",
                        lambda m: f"{m.group('base')}{m.group('exp')}",
                        md_text,
                    )
                    # Case 2: $$ base $$ then $$ exponent $$
                    md_text = _re.sub(
                        r"(?ms)\$\$\n(?P<base>[^\n]+?)\n\$\$\s*\n\$\$\n(?P<exp>[\^_]\{[^}]+\})\n\$\$",
                        lambda m: f"$$\n{m.group('base')}{m.group('exp')}\n$$",
                        md_text,
                    )
                except Exception:
                    pass

                if md_text and str(md_text).strip():
                    page_count = 0
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            page_count = len(pdf.pages)
                    except Exception:
                        page_count = 0

                    images = image_encoder.extract_from_pdf(file_path) if settings.vlm_model else []
                    return LoadedDocument(
                        content=str(md_text),
                        metadata={
                            "source": filename,
                            "type": "pdf",
                            "pages": page_count,
                            "extraction": "marker",
                            "marker_renderer": "chunk",
                            "force_ocr": bool(getattr(settings, "pdf_marker_force_ocr", False)),
                            "outline_headings": [{"level": lvl, "title": title} for (lvl, title) in outline_headings],
                        },
                        images=[{"data": img[0], "page": img[1], "name": f"page_{img[1]}_img"} for img in images],
                    )
            except Exception as e:
                logger.warning(f"Marker PDF->chunk failed for {filename}, falling back to pdfplumber: {e}")

            if mode == "marker":
                # marker 模式下 marker 失败才会继续 fallback
                pass
            else:
                # auto 模式下 marker 失败也会 fallback
                pass

        # Fallback: pdfplumber text extraction
        text_parts = []
        page_count = 0
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {i+1}]\n{page_text}")

                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._table_to_text(table)
                            text_parts.append(f"[Table on Page {i+1}]\n{table_text}")
        except Exception as e:
            logger.error(f"Error loading PDF {filename}: {e}")
            raise

        content = "\n\n".join(text_parts)
          
        # 提取图片
        images = image_encoder.extract_from_pdf(file_path) if settings.vlm_model else []
          
        return LoadedDocument(
            content=content,
            metadata={"source": filename, "type": "pdf", "pages": page_count, "extraction": "pdfplumber"},
            images=[{"data": img[0], "page": img[1], "name": f"page_{img[1]}_img"} for img in images]
        )
    
    def _load_word(self, file_path: str, filename: str) -> LoadedDocument:
        """加载Word文档"""
        logger.info(f"Loading Word: {filename}")
        
        suffix = Path(filename).suffix.lower()
        
        # 处理旧版 .doc 格式
        if suffix == ".doc":
            return self._load_legacy_doc(file_path, filename)
        
        # 处理 .docx 格式
        content = ""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # 处理标题样式
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            paragraphs.append(f"{'#' * int(level)} {para.text}")
                        except ValueError:
                            paragraphs.append(para.text)
                    else:
                        paragraphs.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    paragraphs.append(self._table_to_text(table_data))

            content = "\n\n".join(paragraphs)
                     
        except Exception as e:
            logger.error(f"Error loading Word {filename}: {e}")
            raise
        
        # 提取图片
        images = image_encoder.extract_from_docx(file_path) if settings.vlm_model else []
        
        return LoadedDocument(
            content=content,
            metadata={"source": filename, "type": "word"},
            images=[{"data": img[0], "page": img[1], "name": f"word_img_{i}"} for i, img in enumerate(images)]
        )
    
    def _load_legacy_doc(self, file_path: str, filename: str) -> LoadedDocument:
        """加载旧版 .doc 格式文档（使用 antiword）"""
        import subprocess
        
        try:
            # 尝试使用 antiword 提取文本
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0:
                content = result.stdout
                if content.strip():
                    return LoadedDocument(
                        content=content,
                        metadata={"source": filename, "type": "doc"}
                    )
            
            # antiword 失败，尝试使用 catdoc
            result = subprocess.run(
                ['catdoc', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                return LoadedDocument(
                    content=result.stdout,
                    metadata={"source": filename, "type": "doc"}
                )
            
            # 两种方法都失败
            raise ValueError(
                f"无法解析 .doc 文件 '{filename}'。请尝试将文件另存为 .docx 格式后重新上传。"
            )
            
        except FileNotFoundError:
            raise ValueError(
                f"不支持旧版 .doc 格式文件 '{filename}'。请将文件另存为 .docx 格式后重新上传。"
            )
        except subprocess.TimeoutExpired:
            raise ValueError(f"处理 .doc 文件 '{filename}' 超时")
    
    def _load_excel(self, file_path: str, filename: str) -> LoadedDocument:
        """加载Excel文档"""
        logger.info(f"Loading Excel: {filename}")
        
        try:
            # 读取所有sheet
            xlsx = pd.ExcelFile(file_path)
            all_content = []
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                if not df.empty:
                    all_content.append(f"[Sheet: {sheet_name}]")
                    # 填充空值
                    df = df.fillna('')  # 将NaN替换为空字符串
                    # 将Unnamed列名替换为空字符串
                    df.columns = ['' if str(col).startswith('Unnamed:') else col for col in df.columns]
                    all_content.append(df.to_string(index=False))
                    
        except Exception as e:
            logger.error(f"Error loading Excel {filename}: {e}")
            raise
        
        content = "\n\n".join(all_content)
        return LoadedDocument(
            content=content,
            metadata={"source": filename, "type": "excel", "sheets": len(xlsx.sheet_names)}
        )
    
    def _load_markdown(self, file_path: str, filename: str) -> LoadedDocument:
        """加载Markdown文档"""
        logger.info(f"Loading Markdown: {filename}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            logger.error(f"Error loading Markdown {filename}: {e}")
            raise
        
        return LoadedDocument(
            content=content,
            metadata={"source": filename, "type": "markdown"}
        )
    
    def _load_text(self, file_path: str, filename: str) -> LoadedDocument:
        """加载纯文本文档"""
        logger.info(f"Loading Text: {filename}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()
        
        return LoadedDocument(
            content=content,
            metadata={"source": filename, "type": "text"}
        )
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """将表格转换为文本"""
        if not table:
            return ""
        
        # 使用Markdown表格格式
        lines = []
        for i, row in enumerate(table):
            line = "| " + " | ".join(str(cell) if cell else "" for cell in row) + " |"
            lines.append(line)
            if i == 0:
                # 添加表头分隔符
                separator = "| " + " | ".join("---" for _ in row) + " |"
                lines.append(separator)
        
        return "\n".join(lines)


# 单例
document_loader = DocumentLoader()
